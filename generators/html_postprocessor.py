
import re
from pathlib import Path
from config import UserConfig
from core.logger import get_logger

logger = get_logger("html_postprocessor")

class HTMLPostProcessor:
    
    PARAGRAPH_SPACING_CSS = (Path(__file__).parent / "html_postprocessor_styles.css").read_text(encoding="utf-8")
    
    def __init__(self, html_file_path: str):
        self.html_path = Path(html_file_path)
        self.output_dir = self.html_path.parent
        
        current = self.html_path.parent
        html_base = None
        while current.parent != current:
            if current.name == 'html' and current.parent.name == 'output':
                html_base = current
                output_base = current.parent
                break
            current = current.parent
        
        if html_base:
            try:
                relative_to_html = self.html_path.parent.relative_to(html_base)
                
                if str(relative_to_html) != '.':
                    self.pdf_dir = output_base / 'pdf' / relative_to_html
                    self.docx_dir = output_base / 'docx' / relative_to_html
                else:
                    self.pdf_dir = output_base / 'pdf'
                    self.docx_dir = output_base / 'docx'
            except ValueError:
                self.pdf_dir = Path("output/pdf")
                self.docx_dir = Path("output/docx")
        else:
            self.pdf_dir = Path("output/pdf")
            self.docx_dir = Path("output/docx")
        
        self.pdf_dir.mkdir(parents=True, exist_ok=True)
        self.docx_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"初始化HTML后处理器: {self.html_path.name}")
    
    def check_paragraph_spacing_exists(self, html_content: str) -> bool:
        if 'CSS Version: 2.3' in html_content:
            return True
        
        required_patterns = [
            r'\.article-content\s+p\s*\{[^}]*margin-bottom',
            r'\.article-content\s+blockquote\s*\{[^}]*margin-bottom',
            r'\.article-content\s+figure\s*\{[^}]*margin-bottom',
            r'blockquote\s*\+\s*p\s*\{',
            r'figure\s*\+\s*p\s*\{',
        ]
        
        return all(re.search(pattern, html_content) for pattern in required_patterns)
    
    def fix_paragraph_spacing(self, html_content: str) -> str:
        if self.check_paragraph_spacing_exists(html_content):
            logger.info("内容样式CSS已存在，跳过")
            return html_content
        
        style_end_pattern = r'(\s*)(</style>)'
        match = re.search(style_end_pattern, html_content)
        
        if not match:
            logger.warning("未找到</style>标签，跳过内容样式优化")
            return html_content
        
        indent = match.group(1)
        new_css = f"\n{indent}{self.PARAGRAPH_SPACING_CSS.strip()}\n{indent}"
        fixed_html = html_content[:match.start()] + new_css + match.group(2) + html_content[match.end():]
        
        logger.info("✅ 内容样式CSS已添加（段落/列表/标题/表格美化）")
        return fixed_html
    
    def filter_sensitive_words(self, html_content: str) -> str:
        from bs4 import BeautifulSoup
        
        sensitive_words = UserConfig.SENSITIVE_WORDS
        if not sensitive_words:
            logger.info("未配置敏感词，跳过过滤")
            return html_content
        
        soup = BeautifulSoup(html_content, 'html.parser')
        filtered_count = 0
        
        for text_node in soup.find_all(string=True):
            if text_node.parent.name in ['script', 'style']:
                continue
            
            filtered_text = str(text_node)
            original_text = filtered_text
            
            for word in sensitive_words:
                pattern = re.compile(re.escape(word), re.IGNORECASE)
                filtered_text = pattern.sub('***', filtered_text)
            
            if filtered_text != original_text:
                text_node.replace_with(filtered_text)
                filtered_count += 1
        
        logger.info(f"✅ 敏感词过滤完成，处理了 {filtered_count} 个文本节点")
        return str(soup)
    
    def clean_heading_and_block_spacing(self, html_content: str) -> str:
        from bs4 import BeautifulSoup, NavigableString
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        target_tags = ['h3', 'h4', 'h5', 'blockquote', 'ul', 'ol', 'figure']
        
        total_cleaned = 0
        tag_stats = {}
        
        for tag_name in target_tags:
            tag_stats[tag_name] = 0
            
            for element in soup.find_all(tag_name):
                cleaned_this_element = False
                
                next_elem = element.next_sibling
                removed_count = 0
                
                while next_elem and removed_count < 5:
                    should_remove = False
                    
                    if next_elem.name == 'br':
                        should_remove = True
                    
                    elif isinstance(next_elem, NavigableString):
                        if not next_elem.strip():
                            should_remove = True
                    
                    elif next_elem.name in ['p', 'div']:
                        text_content = next_elem.get_text(strip=True)
                        if not text_content:
                            should_remove = True
                    
                    if should_remove:
                        temp = next_elem.next_sibling
                        next_elem.extract()
                        next_elem = temp
                        removed_count += 1
                        cleaned_this_element = True
                    else:
                        break
                
                if tag_name in ['h3', 'h4', 'h5']:
                    prev_elem = element.previous_sibling
                    prev_br_count = 0
                    
                    temp_elem = prev_elem
                    while temp_elem:
                        if temp_elem.name == 'br':
                            prev_br_count += 1
                            temp_elem = temp_elem.previous_sibling
                        elif isinstance(temp_elem, NavigableString) and not temp_elem.strip():
                            temp_elem = temp_elem.previous_sibling
                        else:
                            break
                    
                    if prev_br_count > 2:
                        removed = 0
                        while prev_elem and removed < (prev_br_count - 1):
                            if prev_elem.name == 'br':
                                temp = prev_elem.previous_sibling
                                prev_elem.extract()
                                prev_elem = temp
                                removed += 1
                                cleaned_this_element = True
                            elif isinstance(prev_elem, NavigableString) and not prev_elem.strip():
                                temp = prev_elem.previous_sibling
                                prev_elem.extract()
                                prev_elem = temp
                            else:
                                break
                
                if cleaned_this_element:
                    tag_stats[tag_name] += 1
                    total_cleaned += 1
        
        if total_cleaned > 0:
            logger.info(f"✅ 清理了 {total_cleaned} 个标签的多余间距")
            for tag, count in tag_stats.items():
                if count > 0:
                    logger.info(f"   - {tag}: {count} 个")
        else:
            logger.info("✅ 未发现需要清理的多余标签")
        
        return str(soup)
    
    def remove_advertisement_articles(self, html_content: str) -> tuple:
        from bs4 import BeautifulSoup

        ad_keywords = UserConfig.AD_KEYWORDS
        if not ad_keywords:
            logger.info("未配置广告关键词，跳过广告去除")
            return html_content, 0

        soup = BeautifulSoup(html_content, 'html.parser')
        removed_count = 0
        removed_titles = []
        removed_ad_numbers = []

        articles = soup.find_all('div', class_='article')

        for article in articles:
            title_tag = article.find('h2')
            if not title_tag:
                continue

            title_text = title_tag.get_text(strip=True)

            has_not_in_toc_attr = article.get('data-not-in-toc') == 'true'

            keyword_match = any(keyword.lower() in title_text.lower() for keyword in ad_keywords)

            is_ad = has_not_in_toc_attr or keyword_match

            if is_ad:
                article_id = article.get('id', '')
                if article_id.startswith('article-'):
                    article_num = article_id.replace('article-', '')
                    removed_ad_numbers.append(article_num)

                article.decompose()
                removed_count += 1
                removed_titles.append(title_text)

        if removed_ad_numbers:
            sidebar_toc = soup.find('nav', class_='table-of-contents')
            if sidebar_toc:
                toc_items = sidebar_toc.find_all('li', class_='toc-item')
                for item in toc_items:
                    link = item.find('a')
                    if link:
                        href = link.get('href', '')
                        for num in removed_ad_numbers:
                            if href == f'#article-{num}':
                                item.decompose()
                                break

                if removed_count > 0:
                    remaining_items = sidebar_toc.find_all('li', class_='toc-item')
                    for new_idx, item in enumerate(remaining_items, start=1):
                        link = item.find('a')
                        if link:
                            link['href'] = f'#article-{new_idx}'

                        toc_number_span = item.find('span', class_='toc-number')
                        if toc_number_span:
                            current_text = toc_number_span.get_text(strip=True)
                            has_marker = '📌' in current_text
                            marker = '📌 ' if has_marker else ''
                            toc_number_span.string = f'{marker}{new_idx}'

        if removed_count > 0:
            remaining_articles = soup.find_all('div', class_='article')
            for new_idx, article in enumerate(remaining_articles, start=1):
                article['id'] = f'article-{new_idx}'

                article_number_div = article.find('div', class_='article-number')
                if article_number_div:
                    article_number_div.string = f'Article / 文章 {new_idx}'

        if removed_ad_numbers:
            toc_table = soup.find('div', class_='table-of-contents-page')
            if toc_table:
                table = toc_table.find('table')
                if table:
                    tbody = table.find('tbody')
                    if tbody:
                        rows = tbody.find_all('tr')
                        for row in rows:
                            first_td = row.find('td')
                            if first_td:
                                td_text = first_td.get_text(strip=True)
                                article_num_text = td_text.replace('📌', '').strip()
                                if article_num_text in removed_ad_numbers:
                                    row.decompose()

                        if tbody and removed_count > 0:
                            remaining_rows = tbody.find_all('tr')
                            for new_idx, row in enumerate(remaining_rows, start=1):
                                first_td = row.find('td')
                                if first_td:
                                    current_text = first_td.get_text(strip=True)
                                    has_marker = '📌' in current_text
                                    marker = '📌' if has_marker else ''
                                    first_td.string = f'{new_idx}{marker}'

        if removed_count > 0:
            sidebar_header = soup.find('div', class_='sidebar-header')
            if sidebar_header:
                subtitle = sidebar_header.find('div', class_='subtitle')
                if subtitle:
                    subtitle_text = subtitle.get_text()
                    import re

                    total_match = re.search(r'Total / 共:\s*(\d+)\s*articles', subtitle_text)
                    not_in_toc_match = re.search(r'Not in TOC.*?📌(\d+)', subtitle_text)

                    if total_match:
                        old_total = int(total_match.group(1))
                        new_total = old_total - removed_count
                        subtitle_text = re.sub(
                            r'(Total / 共:\s*)\d+(\s*articles)',
                            f'\\g<1>{new_total}\\g<2>',
                            subtitle_text
                        )

                    if not_in_toc_match:
                        old_not_in_toc = int(not_in_toc_match.group(1))
                        new_not_in_toc = max(0, old_not_in_toc - removed_count)
                        subtitle_text = re.sub(
                            r'(Not in TOC.*?📌)\d+',
                            f'\\g<1>{new_not_in_toc}',
                            subtitle_text
                        )

                    subtitle.string = subtitle_text

        if removed_count > 0:
            toc_page = soup.find('div', class_='table-of-contents-page')
            if toc_page:
                stat_div = toc_page.find('div', style=lambda x: x and 'text-align: center' in x and 'font-size: 1.05em' in x)
                if stat_div:
                    stat_html = str(stat_div)

                    stat_html = re.sub(
                        r'(<strong>Total / 总计：</strong>)(\d+)',
                        lambda m: f'{m.group(1)}{int(m.group(2)) - removed_count}',
                        stat_html
                    )

                    stat_html = re.sub(
                        r'(<strong>Not in TOC / 目录外发现：</strong>📌)(\d+)',
                        lambda m: f'{m.group(1)}{max(0, int(m.group(2)) - removed_count)}',
                        stat_html
                    )

                    new_stat_div = BeautifulSoup(stat_html, 'html.parser').div
                    stat_div.replace_with(new_stat_div)

        if removed_count > 0:
            logger.info(f"✅ 广告去除完成，删除了 {removed_count} 篇广告文章")
            for idx, title in enumerate(removed_titles, 1):
                logger.info(f"   {idx}. {title[:80]}...")
        else:
            logger.info("✅ 未检测到广告文章")

        return str(soup), removed_count
    
    def preprocess_html_for_docx(self, html_content: str) -> str:
        html_content = re.sub(r'<aside[^>]*>.*?</aside>', '', html_content, flags=re.DOTALL)
        
        html_content = re.sub(r'<button[^>]*class="menu-toggle"[^>]*>.*?</button>', '', html_content, flags=re.DOTALL)
        
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL)
        
        html_content = html_content.replace('class="main-content"', 'class="main-content" style="margin-left: 0;"')
        
        return html_content
    
    def set_formal_fonts(self, docx_path):
        try:
            from docx import Document
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn
            
            doc = Document(str(docx_path))
            
            for para in doc.paragraphs:
                is_heading = para.style.name.startswith('Heading')
                
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    
                    rPr = run._element.get_or_add_rPr()
                    for elem in rPr.findall(qn('w:rFonts')):
                        rPr.remove(elem)
                    
                    rFonts = OxmlElement('w:rFonts')
                    
                    if is_heading:
                        rFonts.set(qn('w:ascii'), 'Arial')
                        rFonts.set(qn('w:eastAsia'), '黑体')
                        rFonts.set(qn('w:hAnsi'), 'Arial')
                        rFonts.set(qn('w:cs'), 'Arial')
                    else:
                        rFonts.set(qn('w:ascii'), 'Times New Roman')
                        rFonts.set(qn('w:eastAsia'), '宋体')
                        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                        rFonts.set(qn('w:cs'), 'Times New Roman')
                    
                    rPr.append(rFonts)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                
                                rPr = run._element.get_or_add_rPr()
                                for elem in rPr.findall(qn('w:rFonts')):
                                    rPr.remove(elem)
                                
                                rFonts = OxmlElement('w:rFonts')
                                rFonts.set(qn('w:ascii'), 'Times New Roman')
                                rFonts.set(qn('w:eastAsia'), '宋体')
                                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                                rFonts.set(qn('w:cs'), 'Times New Roman')
                                rPr.append(rFonts)
            
            doc.save(str(docx_path))
            logger.info("✅ DOCX字体设置完成")
            
        except Exception as e:
            logger.warning(f"⚠️  DOCX字体设置失败（文档仍可使用）: {e}")
    
    def generate_pdf(self) -> bool:
        try:
            from playwright.sync_api import sync_playwright
            
            output_path = self.pdf_dir / self.html_path.with_suffix('.pdf').name
            logger.info(f"生成PDF: {output_path.name}")
            
            with sync_playwright() as p:
                launch_options = {
                    'headless': True,
                    'args': [
                        '--no-sandbox',
                        '--disable-dev-shm-usage',
                        '--force-device-scale-factor=1',
                        '--disable-font-subpixel-positioning',
                        '--disable-lcd-text',
                        '--font-render-hinting=none',
                    ]
                }
                browser = p.chromium.launch(**launch_options)
                
                page = browser.new_page(
                    viewport={'width': 1200, 'height': 1600},
                    device_scale_factor=1.0
                )

                html_file_url = self.html_path.absolute().as_uri()
                page.goto(html_file_url, wait_until="networkidle")

                page.wait_for_load_state("networkidle")
                page.wait_for_load_state("domcontentloaded")

                try:
                    page.wait_for_function("""
                        () => {
                            const images = document.querySelectorAll('img');
                            return Array.from(images).every(img => img.complete);
                        }
                    """, timeout=10000)
                except Exception as e:
                    logger.warning(f"⚠️  部分图片加载超时: {str(e)[:50]}")

                page.pdf(
                    path=str(output_path),
                    format="A4",
                    print_background=True,
                    display_header_footer=False,
                    outline=True,
                    tagged=True,
                    scale=1.0,
                    prefer_css_page_size=False,
                    margin={
                        "top": "20mm",
                        "right": "15mm",
                        "bottom": "20mm",
                        "left": "15mm"
                    }
                )
                
                browser.close()
            
            logger.info(f"✅ PDF已生成: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ PDF生成失败: {e}")
            return False
    
    def generate_docx(self) -> bool:
        try:
            import pypandoc
            
            output_path = self.docx_dir / self.html_path.with_suffix('.docx').name
            logger.info(f"生成DOCX: {output_path.name}")
            
            html_content = self.html_path.read_text(encoding='utf-8')
            html_content = self.preprocess_html_for_docx(html_content)
            
            extra_args = [
                '--standalone',
                '--toc',
                '--toc-depth=3',
            ]
            
            script_dir = Path(__file__).parent
            reference_docx = script_dir / 'reference.docx'
            if reference_docx.exists():
                extra_args.append(f'--reference-doc={reference_docx}')
            
            pypandoc.convert_text(
                html_content,
                'docx',
                format='html',
                outputfile=str(output_path),
                extra_args=extra_args
            )
            
            self.set_formal_fonts(output_path)
            
            logger.info(f"✅ DOCX已生成: {output_path}")
            return True
            
        except ImportError:
            logger.error("❌ 缺少pypandoc库，跳过DOCX生成")
            return False
        except Exception as e:
            logger.error(f"❌ DOCX生成失败: {e}")
            return False
    
    def process(self):
        logger.info("=" * 80)
        logger.info("开始HTML后处理")
        logger.info("=" * 80)
        
        try:
            html_content = self.html_path.read_text(encoding='utf-8')
        except Exception as e:
            logger.error(f"❌ 读取HTML文件失败: {e}")
            return
        
        modified = False
        
        if UserConfig.HTML_ENABLE_AD_REMOVAL:
            logger.info("🚫 检测并去除广告页...")
            cleaned_content, removed_ad_count = self.remove_advertisement_articles(html_content)
            if removed_ad_count > 0:
                html_content = cleaned_content
                modified = True
                print(f"   ✅ 已删除 {removed_ad_count} 篇广告文章")
        
        logger.info("📝 检查内容样式...")
        fixed_content = self.fix_paragraph_spacing(html_content)
        if fixed_content != html_content:
            html_content = fixed_content
            modified = True
        
        if UserConfig.HTML_ENABLE_SENSITIVE_WORDS_FILTER:
            logger.info("🔍 过滤敏感词...")
            filtered_content = self.filter_sensitive_words(html_content)
            if filtered_content != html_content:
                html_content = filtered_content
                modified = True
        
        if modified:
            try:
                self.html_path.write_text(html_content, encoding='utf-8')
                logger.info("✅ HTML文件已更新")
            except Exception as e:
                logger.error(f"❌ 保存HTML文件失败: {e}")
                return
        
        logger.info("📄 生成PDF...")
        self.generate_pdf()
        
        logger.info("📝 生成DOCX...")
        self.generate_docx()
        
        logger.info("=" * 80)
        logger.info("HTML后处理完成")
        logger.info("=" * 80)

